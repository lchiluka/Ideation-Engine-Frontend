# ---------------------------------------------------------------------------
# utils/docx_export.py  –  polished multi-page Word report builder
# ---------------------------------------------------------------------------
"""Turns Proposal-Writer narrative JSON into a well-formatted DOCX file.

Key features
────────────
* Cover page with timestamp
* 0.75-inch margins
* One concept per section (Heading 1 title)
* Ordered subsections (Heading 2) follow the master `SECTION_ORDER`
* Dict values become 2-column tables; lists/strings become bullet or
  numbered lists depending on the key.
* Optional **`on_each_narrative`** callback lets the caller inject side-effects
  (e.g., saving drafts into the ProposalEditor).

Typical usage
─────────────
```
from utils import docx_export

buf = BytesIO()
docx_export.build_docx_report(refined_concepts, buf,
                              on_each_narrative=my_editor.save)
```
"""

from __future__ import annotations

import re, time, json
from io import BytesIO
from typing import List, Dict, Any, Callable

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.trl_assessor import assess_trl
import logging
# PUBLIC EXPORT
__all__ = ["build_docx_report"]

# ---------------------------------------------------------------------------
# Helper – snake-to-words list normaliser
# ---------------------------------------------------------------------------

def _as_list(v: Any) -> List[str]:
    if v is None:
        return []
    if isinstance(v, str):
        parts = [p.strip(" •;-") for p in re.split(r"(?:\n+|•|;)+", v) if p.strip()]
        return parts or [v]
    if isinstance(v, (list, tuple)):
        return [str(x) for x in v]
    if isinstance(v, dict):
        return [f"{k}: {v[k]}" for k in v]
    return [str(v)]

# Global section order – edit here if Proposal schema evolves
SECTION_ORDER = [
    "executive_summary", "problem_statement", "concept_overview",
    "technical_details", "performance_targets", "manufacturing_process",
    "cost_feasibility", "risks_mitigations", "sustainability",
    "applications", "experimental_design", "validation_plan", "kpi_table",
    "ip_landscape", "references",
]

# ---------------------------------------------------------------------------
# Core builder
# ---------------------------------------------------------------------------

def build_docx_report(
    refined_concepts: List[Dict[str, Any]],
    out_stream: BytesIO | str,
    *,
    on_each_narrative: Callable[[Dict[str, Any]], None] | None = None,
) -> None:
    """Create a multi-page Word document from Proposal-Writer outputs.

    Parameters
    ----------
    refined_concepts : list of dict
        Records that will be passed through your generate_proposal() logic
        prior to formatting.
    out_stream : BytesIO | str
        Open binary handle or file path ending in .docx.
    on_each_narrative : callable, optional
        If provided, called with each **narrative JSON** before the content
        is written to the DOCX – perfect for saving into ProposalEditor.
    """
    doc = Document()

    # --- Cover page -------------------------------------------------------
    doc.add_heading("Comprehensive Concept Build-out", level=0)
    ts = time.strftime("%d %b %Y – %H:%M")
    doc.add_paragraph(f"Generated {ts}", style="Intense Quote")
    doc.add_page_break()

    # Standard 0.75-inch margins
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(0.75)
        sec.top_margin  = sec.bottom_margin = Inches(0.75)

    # ---------------------------------------------------------------------
    # Loop concepts -------------------------------------------------------
    # ---------------------------------------------------------------------
    from agents import AGENTS  # local import to avoid circular deps

    for rec in refined_concepts:
        narrative = AGENTS["Proposal Writer Agent"].act(json.dumps(rec), "")
        # force the title to stay exactly what we passed in
        # (we assume rec["title"] was always set correctly)
        narrative["title"] = rec.get("title")

        # keep the raw title around too, if you need it later
        narrative["original_title"] = rec.get("original_title", rec["title"])
        if on_each_narrative:
            on_each_narrative(narrative)
        # ── Inject Experimental Design / DOE ────────────────────────────────
        try:
            # build a strict-JSON prompt
            exp_prompt = (
                "You are an expert experimental designer.  "
                "For the concept below, *return ONLY valid JSON* in the form:\n"
                "{\n"
                '  "experimental_design": [\n'
                '    "…design 1…",\n'
                '    "…design 2…",\n'
                '    "…design 3…"\n'
                "  ]\n"
                "}\n\n"
                f"Title: {rec.get('title')}\n"
                f"Description: {rec.get('description')}\n"
                "Each design must name key factors, levels, response variable, and a rough N."
            )

            # get the raw text back
            raw = AGENTS["Proposal Writer Agent"].act("", exp_prompt)

            if isinstance(raw, dict):
                parsed = raw
            else:
                try:
                    parsed = json.loads(raw)
                except json.JSONDecodeError:
                    logging.warning("DOE JSON parse failed, falling back to lines.  Raw:\n%s", raw)
                    # fallback: split on lines, strip bullets/numbers
                    lines = [l.strip().lstrip("●-0123456789. ") for l in raw.splitlines() if l.strip()]
                    narrative["experimental_design"] = lines[:3]
                    parsed = {}

            # if we got a parsed dict, grab the list
            if isinstance(parsed, dict):
                narrative["experimental_design"] = parsed.get("experimental_design", [])
        except Exception as e:
             logging.warning("Failed to generate experimental_design at all: %s", e)
        # ── Inject TRL‐assessor references ─────────────────────────
        try:
            # run sync assessor on the concept description
            trl_res, evidence_list = assess_trl(rec.get("description") or rec.get("title"))
            # evidence_list is a list of dicts with 'source_url', 'title', 'snippet'
            # take up to 15–20 of them
            refs = []
            for ev in evidence_list[:20]:
                title = ev.get("title") or ev.get("snippet", "")[:60] + "…"
                url   = ev.get("source_url")
                if url:
                    refs.append(f"{title}: {url}")
            if refs:
                narrative["references"] = refs
        except Exception as e:
            # if something goes wrong, leave whatever references were there
            logging.warning("Failed to fetch TRL references: %s", e)
        # Title
        doc.add_heading(narrative["title"], level=1)

        # Ordered subsections
        for key in SECTION_ORDER:
            if key not in narrative:
                continue
            doc.add_heading(key.replace("_", " ").title(), level=2)
            val = narrative[key]

            # dictionaries → 2-col table
            if isinstance(val, dict):
                tbl = doc.add_table(rows=len(val)+1, cols=2)
                tbl.style = "Table Grid"
                tbl.rows[0].cells[0].text = "Key"; tbl.rows[0].cells[1].text = "Value"
                for row, (k, v) in zip(tbl.rows[1:], val.items()):
                    row.cells[0].text = str(k)
                    row.cells[1].text = str(v)
                continue

            # special case: cost_feasibility dict → TRL highlight + details
            if key == "cost_feasibility" and isinstance(val, dict):
                cf = val
                trl_line = f"TRL {cf.get('trl', '?')} – {cf.get('trl_rationale', '')}"
                doc.add_paragraph(trl_line, style="Intense Quote")
                if cf.get("trl_citations"):
                    for cite in _as_list(cf["trl_citations"]):
                        doc.add_paragraph(cite, style="List Bullet")
                if cf.get("cost_breakdown"):
                    doc.add_paragraph(cf["cost_breakdown"])
                if cf.get("capex_estimate"):
                    doc.add_paragraph(f"CapEx estimate: {cf['capex_estimate']}")
                continue

            # list / string fall-through
            bullet_style = "List Number" if key in ("work_plan", "validation_plan") else "List Bullet"
            for item in _as_list(val):
                doc.add_paragraph(item, style=bullet_style)

        doc.add_page_break()

    # Footer
    footer = doc.sections[-1].footer.paragraphs[0]
    footer.text = "Generated by Agentic Ideation Studio"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)

    # Save
    if hasattr(out_stream, "write"):
        doc.save(out_stream)
    else:
        doc.save(out_stream)

# End of file